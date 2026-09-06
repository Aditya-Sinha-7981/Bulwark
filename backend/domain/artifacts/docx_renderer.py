"""
DOCX artifact renderer (Task 14.a).

Deterministically renders a Word document from validated structured input,
per docs/artifacts.md "Templates" (DOCX) and docs/capabilities.md#create_docx.
The model supplies content only; this module owns every formatting decision
(ADR-10, AGENTS.md §6 rule 6) — there is no per-invocation styling.
"""
from __future__ import annotations

import os
import re
import tempfile
import unicodedata
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt

from backend.repositories.artifacts import create_artifact
from backend.utils import paths as paths_module

_MAX_FILENAME_STEM_LEN = 80


class DocxRenderError(Exception):
    """Raised on any rendering/write/persistence failure.

    Callers (the create_docx capability executor) convert this into a failed
    CapabilityExecution result. No partial file is ever left in
    data/artifacts/ and no Artifact row is created when this is raised
    (docs/artifacts.md "Failure handling").
    """


def _sanitize_filename_stem(title: str) -> str:
    """Derive a readable, filesystem/HTTP-safe stem from `title`.

    docs/artifacts.md "Naming": the on-disk name is `{artifact_id}.docx`
    (collision-proof by construction); this stem is only for the
    user-facing `filename` / Content-Disposition header.
    """
    normalized = unicodedata.normalize("NFKD", title).encode("ascii", "ignore").decode("ascii")
    stem = re.sub(r"[^A-Za-z0-9]+", "_", normalized).strip("_")
    if not stem:
        stem = "Document"
    return stem[:_MAX_FILENAME_STEM_LEN]


def _build_document(payload: dict[str, Any]) -> Document:
    """Populate the fixed DOCX template with payload content.

    Template (docs/artifacts.md "Templates" DOCX): title, a metadata block
    (prepared_by, date), then one heading+body per sections[] entry.
    Formatting choices here (heading levels, italic/size on the metadata
    line) are template code — the single, simple, professional template
    called for in artifacts.md; not a per-request decision.
    """
    document = Document()

    title = payload["title"]
    metadata = payload["metadata"]
    sections = payload["sections"]

    document.add_heading(title, level=0)

    meta_paragraph = document.add_paragraph()
    meta_run = meta_paragraph.add_run(
        f"Prepared by: {metadata['prepared_by']}    |    Date: {metadata['date']}"
    )
    meta_run.italic = True
    meta_run.font.size = Pt(10)

    document.add_paragraph()  # spacer between metadata block and sections

    for section in sections:
        document.add_heading(section["heading"], level=1)
        document.add_paragraph(section["body"])

    return document


def render_docx(artifact_id: str, job_id: str, payload: dict[str, Any]) -> dict[str, Any]:
    """
    Render a validated create_docx payload to disk and persist the Artifact row.

    Args:
        artifact_id: pre-generated UUID this file will be stored under.
        job_id: the Job this artifact belongs to (required — data-model.md
            Artifact.job_id is not nullable, and audit.py's emit() requires
            job_id for every event type except network_check).
        payload: validated input matching docs/capabilities.md#create_docx
            exactly: {title, sections: [{heading, body}], metadata:
            {prepared_by, date}}. Caller (create_docx.py) is responsible
            for validation — this function assumes the shape is already
            correct and does not re-validate it.

    Returns:
        {"artifact_id": str, "filename": str} — the capability's output
        schema (docs/capabilities.md#create_docx).

    Raises:
        DocxRenderError: on any python-docx failure, disk-write failure, or
            Artifact-row persistence failure. Atomic write (docs/artifacts.md
            "Failure handling"): the document is built and saved to a temp
            file in the same directory as the final destination, then moved
            into place with os.replace (atomic on both POSIX and Windows for
            same-volume renames) only on full success. If the Artifact row
            insert fails *after* the file is already in place, the
            orphaned file is removed so a broken Artifact-less file is never
            left in data/artifacts/.
    """
    final_path = paths_module.artifacts_path(artifact_id, ".docx")

    tmp_fd, tmp_path_str = tempfile.mkstemp(suffix=".docx.tmp", dir=paths_module.ARTIFACTS_ROOT)
    os.close(tmp_fd)
    tmp_path = Path(tmp_path_str)

    try:
        document = _build_document(payload)
        document.save(tmp_path)
        size_bytes = tmp_path.stat().st_size
        os.replace(tmp_path, final_path)
    except Exception as exc:
        tmp_path.unlink(missing_ok=True)
        raise DocxRenderError(f"failed to render/write docx: {exc}") from exc

    stem = _sanitize_filename_stem(payload["title"])
    filename = f"{stem}.docx"

    try:
        create_artifact(
            job_id=job_id,
            type="docx",
            filename=filename,
            storage_path=final_path.name,  # relative path under data/artifacts/
            size_bytes=size_bytes,
        )
    except Exception as exc:
        # File exists but its Artifact row does not — remove the orphan
        # rather than leave a file no Artifact row points to.
        final_path.unlink(missing_ok=True)
        raise DocxRenderError(f"failed to persist Artifact row: {exc}") from exc

    return {"artifact_id": artifact_id, "filename": filename}